from flask import Flask, request, jsonify, render_template, send_from_directory
from werkzeug.utils import secure_filename
import os
import json
import http.client
import ssl
from utils.ocr_service import recognize_pdf_text
import threading
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Inches
import re

app = Flask(__name__)

# 配置
UPLOAD_FOLDER = 'uploads'
RESULTS_FOLDER = 'results'
DOWNLOADS_FOLDER = 'downloads'
OCR_RESULTS_FOLDER = 'ocr_results'
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'json'}

# 确保文件夹存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(RESULTS_FOLDER, exist_ok=True)
os.makedirs(DOWNLOADS_FOLDER, exist_ok=True)
os.makedirs(OCR_RESULTS_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULTS_FOLDER'] = RESULTS_FOLDER
app.config['DOWNLOADS_FOLDER'] = DOWNLOADS_FOLDER
app.config['OCR_RESULTS_FOLDER'] = OCR_RESULTS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # max file size

# 存储处理状态和结果
processing_status = {}
pdf_results = {}
analysis_results = {}  # 新增：存储分析结果


def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_ocr_results(task_id, pdf_results_data):
    """保存OCR识别结果为JSON文件"""
    try:
        # 构建保存的数据结构
        save_data = {
            'task_id': task_id,
            'timestamp': datetime.now().isoformat(),
            'ocr_results': pdf_results_data
        }
        
        # 生成文件名
        filename = f"ocr_results_{task_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        file_path = os.path.join(app.config['OCR_RESULTS_FOLDER'], filename)
        
        # 保存为JSON文件
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2)
        
        return filename, file_path
        
    except Exception as e:
        print(f"保存OCR结果失败: {e}")
        return None, None


def load_text_file_content(file_path):
    """加载TXT或JSON文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            if file_path.endswith('.json'):
                data = json.load(f)
                # 如果是OCR结果格式
                if 'ocr_results' in data:
                    return data['ocr_results']
                # 如果是页面数组格式
                elif isinstance(data, list):
                    return [{'pages': data, 'filename': os.path.basename(file_path)}]
                else:
                    return data
            else:  # TXT文件
                content = f.read()
                # 将TXT内容按双换行符分割为页面
                pages = [page.strip() for page in content.split('\n\n') if page.strip()]
                return [{'pages': pages, 'filename': os.path.basename(file_path)}]
    except Exception as e:
        print(f"加载文本文件失败: {e}")
        return None


def generate_word_document(content, task_id):
    """将分析结果生成Word文档"""
    try:
        # 创建新文档
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading('法律意见书', 0)
        title.alignment = 1  # 居中对齐
        
        # 添加生成时间
        time_paragraph = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        time_paragraph.alignment = 1  # 居中对齐
        
        # 添加空行
        doc.add_paragraph('')
        
        # 查找"最终法律意见书⬇️"标记，只处理该标记之后的内容
        opinion_content = content
        markers = ["最终法律意见书⬇️", "最终法律意见书⬇", "最终法律意见书↓", "最终法律意见书"]
        
        for marker in markers:
            if marker in content:
                # 找到标记位置，提取标记之后的内容
                marker_index = content.find(marker)
                if marker_index != -1:
                    # 从标记后开始提取内容
                    opinion_content = content[marker_index + len(marker):].strip()
                    print(f"找到标记: {marker}，提取意见书内容")
                    break
        
        if opinion_content == content:
            print("未找到'最终法律意见书'标记，使用全部内容")
        
        # 处理内容，将content按行分割并格式化
        lines = opinion_content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    current_paragraph = None
                continue
                
            # 检查是否是标题（包含#号或特定关键词）
            if (line.startswith('#') or 
                '意见书' in line or 
                '专项法律' in line or
                line.endswith('：') or
                '风险分析' in line or
                '应对策略' in line or
                '结论' in line or
                '建议' in line):
                # 添加标题
                if line.startswith('#'):
                    line = line.lstrip('#').strip()
                heading_level = 1 if ('意见书' in line or '专项法律' in line) else 2
                doc.add_heading(line, heading_level)
                current_paragraph = None
            else:
                # 添加正文段落
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph(line)
                else:
                    current_paragraph.add_run('\n' + line)
        
        # 保存文档
        filename = f"法律意见书_{task_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(app.config['DOWNLOADS_FOLDER'], filename)
        doc.save(file_path)
        
        return filename, file_path
        
    except Exception as e:
        print(f"生成Word文档失败: {e}")
        return None, None


# def process_pdfs_async(task_id, pdf_files):
#     """异步处理PDF文件"""
#     try:
#         processing_status[task_id] = {'status': 'processing', 'progress': 0, 'total': len(pdf_files)}
#         pdf_results[task_id] = []
        
#         for i, pdf_file in enumerate(pdf_files):
#             try:
#                 # 处理单个PDF
#                 pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_file)
#                 page_texts = recognize_pdf_text(pdf_path)
                
#                 # 保存结果
#                 result = {
#                     'filename': pdf_file,
#                     'pages': page_texts,
#                     'timestamp': datetime.now().isoformat()
#                 }
#                 pdf_results[task_id].append(result)
                
#                 # 更新进度
#                 processing_status[task_id]['progress'] = i + 1
                
#             except Exception as e:
#                 # 记录错误但继续处理其他文件
#                 error_result = {
#                     'filename': pdf_file,
#                     'error': str(e),
#                     'timestamp': datetime.now().isoformat()
#                 }
#                 pdf_results[task_id].append(error_result)
#                 processing_status[task_id]['progress'] = i + 1
        
#         processing_status[task_id]['status'] = 'completed'
        
#     except Exception as e:
#         processing_status[task_id] = {'status': 'error', 'error': str(e)}


def process_pdfs_async(task_id, pdf_files):
    """异步处理PDF文件"""
    try:
        processing_status[task_id] = {'status': 'processing', 'progress': 0, 'total': len(pdf_files)}
        pdf_results[task_id] = []
        
        for i, pdf_file in enumerate(pdf_files):
            try:
                # 处理单个PDF
                pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_file)
                page_texts = recognize_pdf_text(pdf_path)
                
                # 保存结果
                result = {
                    'filename': pdf_file,
                    'pages': page_texts,
                    'timestamp': datetime.now().isoformat()
                }
                pdf_results[task_id].append(result)
                
                # OCR处理完成后删除原始PDF文件
                try:
                    os.remove(pdf_path)
                    print(f"已删除处理完成的文件: {pdf_file}")
                except Exception as e:
                    print(f"删除文件失败: {pdf_file}, 错误: {e}")
                
                # 更新进度
                processing_status[task_id]['progress'] = i + 1
                
            except Exception as e:
                # 记录错误但继续处理其他文件
                error_result = {
                    'filename': pdf_file,
                    'error': str(e),
                    'timestamp': datetime.now().isoformat()
                }
                pdf_results[task_id].append(error_result)
                processing_status[task_id]['progress'] = i + 1
                
                # 即使处理失败也删除文件
                try:
                    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_file)
                    os.remove(pdf_path)
                except:
                    pass
        
        processing_status[task_id]['status'] = 'completed'
        
        # 自动保存OCR结果
        try:
            ocr_filename, ocr_filepath = save_ocr_results(task_id, pdf_results[task_id])
            if ocr_filename:
                processing_status[task_id]['ocr_file'] = ocr_filename
                print(f"OCR结果已自动保存: {ocr_filename}")
        except Exception as e:
            print(f"自动保存OCR结果失败: {e}")
        
    except Exception as e:
        processing_status[task_id] = {'status': 'error', 'error': str(e)}

def call_workflow_api(pdf_content_list):
    print("pdf_content_list:")
    print(pdf_content_list)
    """调用workflow API"""
    headers = {
        "Content-Type": "application/json",
        "Accept": "text/event-stream",
        "Authorization": "Bearer 98fb62695047338d32729257a65a48a6:NDNkOTVmZmJjYzc0OTg5MTg5ODI5MDNi",
    }
    
    data = {
        "flow_id": "7358700018745618434",
        "uid": "123",
        "parameters": {
            "AGENT_USER_INPUT": "请分析以下PDF内容",
            "pdf_list": pdf_content_list
        },
        "ext": {"bot_id": "adjfidjf", "caller": "workflow"},
        "stream": True,
    }
    
    payload = json.dumps(data)
    
    try:
        conn = http.client.HTTPSConnection("xingchen-api.xf-yun.com", timeout=300)
        conn.request("POST", "/workflow/v1/chat/completions", payload, headers, encode_chunked=True)
        res = conn.getresponse()
        
        response_content = []
        
        if data.get("stream"):
            while chunk := res.readline():
                chunk_str = chunk.decode("utf-8").strip()
                if chunk_str.startswith("data: "):
                    json_str = chunk_str[6:]
                    try:
                        data_dict = json.loads(json_str)
                        if "choices" in data_dict and len(data_dict["choices"]) > 0:
                            choice = data_dict["choices"][0]
                            if "delta" in choice and "content" in choice["delta"]:
                                content = choice["delta"]["content"]
                                print("content:",content)
                                response_content.append(content)
                    except json.JSONDecodeError:
                        continue
        
        return ''.join(response_content)
        
    except Exception as e:
        return f"API调用错误: {str(e)}"


def call_workflow_api_stream(pdf_content_list):
    """流式调用workflow API"""
    print("pdf_content_list:")
    print(pdf_content_list)
    
    headers = {
        "Content-Type": "application/json",
        "Accept": "text/event-stream",
        "Authorization": "Bearer 98fb62695047338d32729257a65a48a6:NDNkOTVmZmJjYzc0OTg5MTg5ODI5MDNi",
    }
    
    data = {
        "flow_id": "7358700018745618434",
        "uid": "123",
        "parameters": {
            "AGENT_USER_INPUT": "请分析以下PDF内容",
            "pdf_list": pdf_content_list
        },
        "ext": {"bot_id": "adjfidjf", "caller": "workflow"},
        "stream": True,
    }
    
    payload = json.dumps(data)
    
    try:
        conn = http.client.HTTPSConnection("xingchen-api.xf-yun.com", timeout=320)
        conn.request("POST", "/workflow/v1/chat/completions", payload, headers, encode_chunked=True)
        res = conn.getresponse()
        
        if data.get("stream"):
            while chunk := res.readline():
                chunk_str = chunk.decode("utf-8").strip()
                if chunk_str.startswith("data: "):
                    json_str = chunk_str[6:]
                    try:
                        data_dict = json.loads(json_str)
                        if "choices" in data_dict and len(data_dict["choices"]) > 0:
                            choice = data_dict["choices"][0]
                            if "delta" in choice and "content" in choice["delta"]:
                                content = choice["delta"]["content"]
                                if content.startswith("正在"):
                                    content = f"\n【系统提示】{content}\n"
                                print("content:",content)
                                yield content
                    except json.JSONDecodeError:
                        continue
        
        conn.close()
        
    except Exception as e:
        yield f"API调用错误: {str(e)}"


@app.route('/')
def index():
    """PDF分析页面"""
    return render_template('analysis.html')


@app.route('/dashboard')
def dashboard():
    """系统主控台页面"""
    return render_template('dashboard.html')


@app.route('/original')
def original_index():
    """原始分析页面"""
    return render_template('index.html')


@app.route('/static/logo.jpg')
def serve_logo():
    """提供logo图片"""
    
    return send_from_directory('.', 'static/logo.jpg')


@app.route('/upload', methods=['POST'])
def upload_files():
    """上传PDF/TXT/JSON文件"""
    if 'files' not in request.files:
        return jsonify({'error': '没有选择文件'}), 400
    
    files = request.files.getlist('files')
    if not files or all(file.filename == '' for file in files):
        return jsonify({'error': '没有选择文件'}), 400
    
    # 生成任务ID
    task_id = str(uuid.uuid4())
    uploaded_files = []
    text_files = []
    
    for file in files:
        if file and allowed_file(file.filename):
            filename = file.filename
            # 添加时间戳避免文件名冲突
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{timestamp}_{filename}"
            
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # 检查文件类型
            if filename.lower().endswith(('.txt', '.json')):
                text_files.append(filename)
            else:
                uploaded_files.append(filename)
    
    if not uploaded_files and not text_files:
        return jsonify({'error': '没有有效的文件'}), 400
    
    # 如果有文本文件，直接处理
    if text_files:
        try:
            pdf_results[task_id] = []
            for text_file in text_files:
                text_path = os.path.join(app.config['UPLOAD_FOLDER'], text_file)
                content = load_text_file_content(text_path)
                if content:
                    pdf_results[task_id].extend(content)
                
                # 处理完后删除文本文件
                try:
                    os.remove(text_path)
                except:
                    pass
            
            processing_status[task_id] = {'status': 'completed', 'progress': 1, 'total': 1, 'file_type': 'text'}
            
            return jsonify({
                'task_id': task_id,
                'message': f'已加载 {len(text_files)} 个文本文件',
                'files': text_files,
                'file_type': 'text'
            })
        except Exception as e:
            return jsonify({'error': f'处理文本文件失败: {str(e)}'}), 500
    
    # 如果有PDF文件，启动异步OCR处理
    if uploaded_files:
        thread = threading.Thread(target=process_pdfs_async, args=(task_id, uploaded_files))
        thread.start()
        
        return jsonify({
            'task_id': task_id,
            'message': f'开始处理 {len(uploaded_files)} 个PDF文件',
            'files': uploaded_files,
            'file_type': 'pdf'
        })


@app.route('/status/<task_id>')
def get_status(task_id):
    """获取处理状态"""
    if task_id not in processing_status:
        return jsonify({'error': '任务不存在'}), 404
    
    status = processing_status[task_id]
    
    # 如果处理完成，返回结果
    if status.get('status') == 'completed' and task_id in pdf_results:
        return jsonify({
            'status': 'completed',
            'results': pdf_results[task_id]
        })
    
    return jsonify(status)


@app.route('/analyze', methods=['POST'])
def analyze_pdfs():
    """分析PDF内容"""
    data = request.get_json()
    task_id = data.get('task_id')
    k_pages = data.get('k_pages', 5)  # 每K页合并，默认5页
    
    if not task_id or task_id not in pdf_results:
        return jsonify({'error': '无效的任务ID或结果不存在'}), 400
    
    try:
        # 获取所有PDF的页面内容
        all_pages = []
        for pdf_result in pdf_results[task_id]:
            if 'pages' in pdf_result:
                all_pages.extend(pdf_result['pages'])
        
        if not all_pages:
            return jsonify({'error': '没有找到可分析的内容'}), 400
        
        # 按K页合并内容
        pdf_content_list = []
        for i in range(0, len(all_pages), k_pages):
            chunk = all_pages[i:i + k_pages]
            combined_content = '\n\n'.join(chunk)
            pdf_content_list.append(combined_content)
        
        # 调用workflow API
        api_response = call_workflow_api(pdf_content_list)
        
        # 保存分析结果
        analysis_results[task_id] = {
            'content': api_response,
            'timestamp': datetime.now().isoformat(),
            'chunks_count': len(pdf_content_list),
            'k_pages': k_pages
        }
        
        # 生成Word文档
        doc_filename, doc_filepath = generate_word_document(api_response, task_id)
        
        response_data = {
            'success': True,
            'chunks_count': len(pdf_content_list),
            'k_pages': k_pages,
            'analysis_result': api_response
        }
        
        # 如果Word文档生成成功，添加下载链接
        if doc_filename:
            response_data['word_document'] = {
                'filename': doc_filename,
                'download_url': f'/download/{task_id}'
            }
            analysis_results[task_id]['word_filename'] = doc_filename
        
        return jsonify(response_data)
        
    except Exception as e:
        return jsonify({'error': f'分析失败: {str(e)}'}), 500


@app.route('/analyze_stream', methods=['POST'])
def analyze_pdfs_stream():
    """流式分析PDF内容"""
    from flask import Response
    
    data = request.get_json()
    task_id = data.get('task_id')
    k_pages = data.get('k_pages', 5)  # 每K页合并，默认5页
    
    if not task_id or task_id not in pdf_results:
        return jsonify({'error': '无效的任务ID或结果不存在'}), 400
    
    def generate():
        try:
            # 获取所有PDF的页面内容
            all_pages = []
            for pdf_result in pdf_results[task_id]:
                if 'pages' in pdf_result:
                    all_pages.extend(pdf_result['pages'])
            
            if not all_pages:
                yield f"data: {json.dumps({'error': '没有找到可分析的内容'})}\n\n"
                return
            
            # 按K页合并内容
            pdf_content_list = []
            for i in range(0, len(all_pages), k_pages):
                chunk = all_pages[i:i + k_pages]
                combined_content = '\n\n'.join(chunk)
                pdf_content_list.append(combined_content)
            
            # 发送初始信息
            yield f"data: {json.dumps({'type': 'init', 'chunks_count': len(pdf_content_list), 'k_pages': k_pages})}\n\n"
            
            # 收集完整的分析结果
            full_response = []
            
            # 调用workflow API并流式传输
            for content in call_workflow_api_stream(pdf_content_list):
                if content:
                    full_response.append(content)
                    yield f"data: {json.dumps({'type': 'content', 'data': content})}\n\n"
            
            # 分析完成后生成Word文档
            if full_response:
                complete_response = ''.join(full_response)
                
                # 保存分析结果
                analysis_results[task_id] = {
                    'content': complete_response,
                    'timestamp': datetime.now().isoformat(),
                    'chunks_count': len(pdf_content_list),
                    'k_pages': k_pages
                }
                
                # 生成Word文档
                doc_filename, doc_filepath = generate_word_document(complete_response, task_id)
                
                if doc_filename:
                    analysis_results[task_id]['word_filename'] = doc_filename
                    yield f"data: {json.dumps({'type': 'word_generated', 'filename': doc_filename, 'download_url': f'/download/{task_id}'})}\n\n"
            
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    
    return Response(generate(), mimetype='text/plain', headers={
        'Cache-Control': 'no-cache',
        'Connection': 'keep-alive',
        'X-Accel-Buffering': 'no'
    })


@app.route('/results/<task_id>')
def get_results(task_id):
    """获取处理结果详情"""
    if task_id not in pdf_results:
        return jsonify({'error': '结果不存在'}), 404
    
    return jsonify(pdf_results[task_id])


@app.route('/download/<task_id>')
def download_word(task_id):
    """下载Word文档"""
    if task_id not in analysis_results:
        return jsonify({'error': '分析结果不存在'}), 404
    
    result = analysis_results[task_id]
    if 'word_filename' not in result:
        return jsonify({'error': 'Word文档未生成'}), 404
    
    filename = result['word_filename']
    
    try:
        return send_from_directory(
            app.config['DOWNLOADS_FOLDER'], 
            filename, 
            as_attachment=True,
            download_name=filename
        )
    except FileNotFoundError:
        return jsonify({'error': '文件不存在'}), 404


@app.route('/download_ocr/<task_id>')
def download_ocr_results(task_id):
    """下载OCR识别结果"""
    if task_id not in processing_status:
        return jsonify({'error': '任务不存在'}), 404
    
    status = processing_status[task_id]
    if 'ocr_file' not in status:
        # 如果没有自动保存的文件，尝试生成一个
        if task_id in pdf_results:
            ocr_filename, ocr_filepath = save_ocr_results(task_id, pdf_results[task_id])
            if not ocr_filename:
                return jsonify({'error': 'OCR结果生成失败'}), 500
        else:
            return jsonify({'error': 'OCR结果不存在'}), 404
    else:
        ocr_filename = status['ocr_file']
    
    try:
        return send_from_directory(
            app.config['OCR_RESULTS_FOLDER'], 
            ocr_filename, 
            as_attachment=True,
            download_name=ocr_filename
        )
    except FileNotFoundError:
        return jsonify({'error': 'OCR结果文件不存在'}), 404


@app.route('/analysis_info/<task_id>')
def get_analysis_info(task_id):
    """获取分析结果信息"""
    if task_id not in analysis_results:
        return jsonify({'error': '分析结果不存在'}), 404
    
    result = analysis_results[task_id]
    response_data = {
        'timestamp': result['timestamp'],
        'chunks_count': result['chunks_count'],
        'k_pages': result['k_pages'],
        'has_word_document': 'word_filename' in result
    }
    
    if 'word_filename' in result:
        response_data['word_document'] = {
            'filename': result['word_filename'],
            'download_url': f'/download/{task_id}'
        }
    
    return jsonify(response_data)


@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': '文件太大，请选择小于500MB的文件'}), 413


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5500)
